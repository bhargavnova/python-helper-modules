import PyPDF2
from docx import Document
from pdf2image import convert_from_path
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class PDFtoDOCXConverter:
    def __init__(self, pdf_file, docx_file):
        self.pdf_file = pdf_file
        self.docx_file = docx_file

    def convert(self):
        pdf = PyPDF2.PdfReader(self.pdf_file)
        docx = Document()

        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]

            # Extract text from the PDF page
            text = page.extract_text()

            # Add a new page for each PDF page
            if page_num > 0:
                docx.add_page_break()
            
            # Add the extracted text to the DOCX document
            docx.add_paragraph(text).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Extract images from the PDF page
            images = convert_from_path(self.pdf_file, first_page=page_num + 1, last_page=page_num + 1)

            for image_num, image in enumerate(images):
                # Save the image and add it to the DOCX
                image_path = f"page_{page_num + 1}_img_{image_num + 1}.jpg"
                image.save(image_path)
                docx.add_picture(image_path, width=Inches(6))  # Adjust the image width as needed
                docx.add_paragraph("")  # Adds an empty line after the image

        docx.save(self.docx_file)

if __name__ == "__main__":
    pdf_converter = PDFtoDOCXConverter("input.pdf", "output.docx")
    pdf_converter.convert()
